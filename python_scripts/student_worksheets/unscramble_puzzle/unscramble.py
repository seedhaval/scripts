# -*- coding: utf8 -*-

from docx import Document
from docx.shared import Pt
import random
import os

num_words = 20

with open('word_list.txt') as f:
    words = [x.strip() for x in f.readlines() if len(x.strip())>1]

random.shuffle(words)
word_subset = words[:num_words]

document = Document()
document.add_heading('Unscramble Quiz', 0)
p = document.add_paragraph('Unscramble the below words')
for ln in word_subset:
    blanks = ' __'*len(ln)
    chr_ar = [x for x in ln]
    while ln == ''.join(chr_ar):
        random.shuffle(chr_ar)
    run = p.add_run('\n' + ' '.join(chr_ar) + ' : ' + blanks)
    run.font.size = Pt(20)
    run.font.name = "Bahnschrift Light"

document.save('out.docx')
os.startfile('out.docx')
