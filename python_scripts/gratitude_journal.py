import json
from datetime import datetime
import os
from docx import Document

data_fl = "gratitude_data.json"

def get_choice():
    print("Menu")
    print("1. Make new gratitude entry")
    print("2. Generate report")
    return input("Enter choice : ")

def make_gratitude_entry():
    text = input("Provide gratitude summary : ")
    if os.path.exists(data_fl):
        with open(data_fl) as f:
            data = json.load(f)
    else:
        data = {}
    dt = datetime.now().strftime("%Y-%m-%d")
    if dt in data:
        data[dt].append(text)
    else:
        data[dt] = [text]
    with open(data_fl,'w') as f:
        f.write(json.dumps(data, indent=2))

def generate_report():
    with open(data_fl) as f:
        data = json.load(f)
    document = Document()
    document.add_heading("My Gratitude Journal", 0)
    for dt,entries in sorted(data.items(), reverse=True):
        document.add_heading(dt, level=2)
        for entry in entries:
            document.add_paragraph(entry)
        document.add_paragraph("")
    document.save('out.docx')
    os.startfile("out.docx")

ch = get_choice()
if ch == '1':
    make_gratitude_entry()
elif ch == '2':
    generate_report()
else:
    print("Invalid choice. Enter 1 or 2.")
