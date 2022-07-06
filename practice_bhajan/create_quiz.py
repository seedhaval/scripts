import pathlib
import regex
from docx import Document
from docx.shared import Inches, Pt
import random

hw = '्'
def get_choice():
    ar = list(pathlib.Path('.').glob('text/*.txt'))
    print('Select Shlok / Bhajan by number')
    for i,fl in enumerate(ar):
        print(i+1,fl.stem)
    return ar[int(input('Enter number : '))-1]

def split_word(txt):
    chr = regex.findall(r'\X', txt, regex.U)
    for i,c in enumerate(chr):
        if hw in c:
            chr[i] += chr[i+1]
            chr[i+1] = ''
    return [x for x in chr if x]

def get_unscramble_quiz(w):
    out = []
    chr = [split_word(x) for x in set(w)]
    random.shuffle(chr)
    for c in [x for x in chr if len(x) > 2][:5]:
            random.shuffle(c)
            out.append( ' '.join(c) + ' : ' + '   '.join(['__']*len(c)))
    return out
            
def get_complete_line_quiz(txt):
    out = []
    ln_ar = [x.strip() for x in txt.splitlines() if x.strip()]
    random.shuffle(ln_ar)
    for ln in ln_ar[:5]:
        cur_out = []
        words = ln.split()
        blnk = list(range(len(words)))
        random.shuffle(blnk)
        blnk = blnk[:len(words)//2]
        for i,w in enumerate(words):
            c = split_word(w)
            if i in blnk:
                cur_out.append('___'*len(c))
            else:
                cur_out.append(w)
        out.append('   '.join(cur_out))
    return out

def get_search_word_list(words):
    wrd_set = [x for x in set(words) if len(split_word(x)) > 2]
    random.shuffle(wrd_set)
    return wrd_set[:5]

def get_word_table(words,word_set):
    mx_chars = max([len(split_word(x)) for x in word_set])
    mx_words = len(word_set)
    width = max(mx_chars, mx_words) + 2
    ar = []
    for i in range(width):
        ar.append([])
        for j in range(width):
            ar[i].append('x')
            
    col_ar = list(range(width))
    random.shuffle(col_ar)
    for i,cid in list(enumerate(col_ar[:mx_words])):
        c = split_word(word_set[i])
        dirc = random.choice([0,1])
        if dirc == 0:
            strt = random.randint(0,width-len(c))            
            for j,v in list(enumerate(c)):               
                ar[strt+j][cid] = v
        elif dirc == 1:
            strt = random.randint(width-len(c),width-1)            
            for j,v in list(enumerate(c)):               
                ar[strt-j][cid] = v
    all_char = list(set([x for y in words for x in split_word(y)]))           
    for i in range(len(ar)):
        for j in range(len(ar)):
            if ar[i][j] == 'x':
                ar[i][j] = random.choice(all_char)
                
    return ar
                                                 
ch = get_choice()        
print(ch.stem)
txt = ch.read_text()
words = txt.split()

document = Document()
document.add_heading(ch.stem + ' Quiz', 0)
document.add_heading('Unscramble', level=2)
p = document.add_paragraph('Unscramble the below words')
for ln in get_unscramble_quiz(words):
    run = p.add_run('\n\n'+ln)
    run.font.size = Pt(24)

document.add_heading('Complete Lines', level=2)
p = document.add_paragraph('Complete the below lines')
for ln in get_complete_line_quiz(txt):
    run = p.add_run('\n\n'+ln)
    run.font.size = Pt(24)
    
document.add_heading('Word Search', level=2)
p = document.add_paragraph('Search below words in the table on next page')
word_set = get_search_word_list(words)
for ln in word_set:
    run = p.add_run('\n\n'+ln)
    run.font.size = Pt(12)

document.add_page_break()
rec = get_word_table(words,word_set)
table = document.add_table(rows=0, cols=len(rec))
for row in rec:
    row_cells = table.add_row().cells
    for i,v in enumerate(row):
        row_cells[i].text = v
        
                                    
document.save('out.docx')
